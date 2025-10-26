import streamlit as st
import time
import textstat
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import json
import pandas as pd
import requests
from typing import Dict, List, Optional
import PyPDF2
import anthropic
import os

# --- Page Configuration ---
st.set_page_config(
    page_title="ICF-GPT",
    page_icon="🩺",
    layout="wide",
    initial_sidebar_state="expanded",
)

# --- Custom CSS for Dark Theme ---
st.markdown("""
<style>
    /* Main theme colors */
    :root {
        --primary-color: #000000;  /* Pure Black (for buttons, main accents) */
        --bg-dark: #FFFFFF;       /* Pure White (for the main page background) */
        --bg-darker: #FFFFFF;      /* Pure White (for other background layers) */
        --bg-card: #F0F0F0;       /* Light Grey (This will fill your boxes) */
        --text-primary: #000000;   /* Pure Black (for main text) */
        --text-secondary: #333333; /* Dark Grey (for secondary text/labels for better contrast) */
        --border-color: #C0C0C0;   /* Medium Grey (for visible borders) */
    }
    
    /* Global background */
    .stApp {
        background-color: var(--bg-darker);
    }
    
    /* Hide the top toolbar/header */
    header[data-testid="stHeader"] {
        display: none !important;
    }
    
    /* Hide the main menu button */
    #MainMenu {
        display: none !important;
    }
    
    /* Hide the footer */
    footer {
        display: none !important;
    }
    
    /* Adjust main container padding since header is gone */
    .main .block-container {
        padding-top: 2rem;
    }
    
    /* Sidebar styling */
    [data-testid="stSidebar"] {
        background-color: var(--bg-dark);
        border-right: 1px solid var(--border-color);
    }
    
    [data-testid="stSidebar"] h1, 
    [data-testid="stSidebar"] h2, 
    [data-testid="stSidebar"] h3 {
        color: var(--text-primary);
        font-weight: 600;
    }
    
    /* Main content area */
    .main .block-container {
        padding-top: 2rem;
        max-width: 1400px;
    }
    
    /* Title styling */
    h1 {
        color: var(--text-primary);
        font-size: 2rem;
        font-weight: 700;
        margin-bottom: 0.5rem;
    }
    
    h1 span:first-child {
        color: var(--primary-color);
    }
    
    /* Tab styling */
    .stTabs [data-baseweb="tab-list"] {
        gap: 1rem;
        background-color: var(--bg-card);
        padding: 0.5rem;
        border-radius: 12px;
        border: 1px solid var(--border-color);
    }
    
    .stTabs [data-baseweb="tab"] {
        background-color: transparent;
        border: none;
        color: var(--text-secondary);
        font-weight: 500;
        padding: 0.75rem 1.5rem;
        border-radius: 8px;
    }
    
    .stTabs [aria-selected="true"] {
        background-color: var(--primary-color) !important;
        color: var(--bg-darker) !important;
        border: none;
    }
    
    .stTabs [data-baseweb="tab"]:hover {
        background-color: rgba(45, 212, 191, 0.1);
    }
    
    /* Remove red underline from active tab */
    .stTabs [data-baseweb="tab-highlight"] {
        background-color: transparent !important;
        display: none !important;
    }
    
    .stTabs [data-baseweb="tab-border"] {
        background-color: transparent !important;
        display: none !important;
    }
    
    /* Card/Form styling */
    [data-testid="stForm"] {
        background-color: var(--bg-card);
        border: 1px solid var(--border-color);
        border-radius: 12px;
        padding: 2rem;
    }
    
    /* Button styling - ALL BUTTONS TEAL */
    .stButton > button,
    .stDownloadButton > button,
    .stFormSubmitButton > button {
        background-color: var(--primary-color) !important;
        color: var(--bg-darker) !important;
        border: none !important;
        border-radius: 8px !important;
        padding: 0.75rem 2rem !important;
        font-weight: 600 !important;
        font-size: 1rem !important;
        width: 100% !important;
        transition: all 0.3s ease !important;
    }
    
    .stButton > button:hover,
    .stDownloadButton > button:hover,
    .stFormSubmitButton > button:hover {
        background-color: #25B8A5 !important;
        box-shadow: 0 4px 12px rgba(45, 212, 191, 0.3) !important;
        transform: translateY(-1px);
    }
    
    .stButton > button:active,
    .stDownloadButton > button:active,
    .stFormSubmitButton > button:active {
        transform: translateY(0);
    }
    
    .stButton > button:focus,
    .stDownloadButton > button:focus,
    .stFormSubmitButton > button:focus,
    .stButton > button:focus-visible,
    .stDownloadButton > button:focus-visible,
    .stFormSubmitButton > button:focus-visible {
        outline: none !important;
        box-shadow: none !important;
        border: none !important;
        background-color: var(--primary-color) !important;
    }
    
    /* Form submit buttons - extra teal emphasis */
    button[kind="primaryFormSubmit"],
    button[type="submit"] {
        background-color: var(--primary-color) !important;
        color: var(--bg-darker) !important;
        border: none !important;
    }
    
    button[kind="primaryFormSubmit"]:focus,
    button[type="submit"]:focus,
    button[kind="primaryFormSubmit"]:focus-visible,
    button[type="submit"]:focus-visible {
        outline: none !important;
        box-shadow: none !important;
        border: none !important;
        background-color: var(--primary-color) !important;
    }
    
    /* Remove any Streamlit default focus styles */
    button:focus-visible {
        outline: none !important;
        box-shadow: none !important;
    }
    
    *:focus-visible {
        outline: none !important;
    }
    
    /* Input fields */
    .stTextInput > div > div > input,
    .stSelectbox > div > div > select,
    .stTextArea textarea {
        background-color: var(--bg-darker) !important;
        color: var(--text-primary) !important;
        border: 1px solid var(--border-color) !important;
        border-radius: 8px;
        padding: 0.75rem;
    }
    
    .stTextInput > div > div > input:focus,
    .stSelectbox > div > div > select:focus,
    .stTextArea textarea:focus {
        border-color: var(--text-secondary) !important;
        box-shadow: none !important;
    }
    
    /* File uploader */
    [data-testid="stFileUploader"] {
        background-color: var(--bg-darker);
        border: 2px dashed var(--border-color);
        border-radius: 12px;
        padding: 2rem;
        text-align: center;
    }
    
    [data-testid="stFileUploader"] section {
        border: none;
        background-color: transparent;
    }
    
    [data-testid="stFileUploader"] section > div {
        color: var(--text-secondary);
    }
    
    [data-testid="stFileUploader"] button {
        background-color: var(--primary-color) !important;
        color: var(--bg-darker) !important;
        border: none !important;
        border-radius: 8px;
        padding: 0.5rem 1.5rem;
        font-weight: 600;
    }
    
    /* Info/warning boxes - all grey tones */
    .stAlert {
        background-color: var(--bg-card) !important;
        border-left: 4px solid var(--text-secondary) !important;
        border-radius: 8px;
        color: var(--text-primary) !important;
    }
    
    .stSuccess {
        background-color: var(--bg-card) !important;
        border-left: 4px solid var(--text-secondary) !important;
        color: var(--text-primary) !important;
    }
    
    .stWarning {
        background-color: var(--bg-card) !important;
        border-left: 4px solid var(--text-secondary) !important;
        color: var(--text-primary) !important;
    }
    
    .stError {
        background-color: var(--bg-card) !important;
        border-left: 4px solid var(--text-secondary) !important;
        color: var(--text-primary) !important;
    }
    
    .stInfo {
        background-color: var(--bg-card) !important;
        border-left: 4px solid var(--text-secondary) !important;
        color: var(--text-primary) !important;
    }
    
    /* Remove any default Streamlit colored backgrounds */
    div[data-testid="stNotificationContentSuccess"],
    div[data-testid="stNotificationContentError"],
    div[data-testid="stNotificationContentWarning"],
    div[data-testid="stNotificationContentInfo"] {
        background-color: var(--bg-card) !important;
    }
    
    /* Metrics */
    [data-testid="stMetricValue"] {
        color: var(--primary-color);
        font-size: 2rem;
        font-weight: 700;
    }
    
    /* Subheaders */
    .stSubheader {
        color: var(--text-primary);
        font-size: 1.25rem;
        font-weight: 600;
        margin-top: 1.5rem;
        margin-bottom: 1rem;
    }
    
    /* Template card in sidebar */
    .template-card {
        background-color: var(--bg-card);
        border: 1px solid var(--border-color);
        border-radius: 8px;
        padding: 1rem;
        margin-bottom: 0.75rem;
    }
    
    .template-card-title {
        color: var(--text-primary);
        font-weight: 600;
        font-size: 0.95rem;
        margin-bottom: 0.25rem;
    }
    
    .template-card-subtitle {
        color: var(--text-secondary);
        font-size: 0.8rem;
    }
    
    /* Expander styling - grey tones only */
    .streamlit-expanderHeader {
        background-color: var(--bg-card) !important;
        border: 1px solid var(--border-color) !important;
        border-radius: 8px;
        color: var(--text-primary) !important;
    }
    
    .streamlit-expanderHeader:hover {
        background-color: var(--bg-card) !important;
        border-color: var(--text-secondary) !important;
    }
    
    /* Dataframe styling */
    .stDataFrame {
        border: 1px solid var(--border-color);
        border-radius: 8px;
    }
    
    /* Progress bars - teal only for actual progress */
    .stProgress > div > div {
        background-color: var(--primary-color) !important;
    }
    
    /* Spinner - teal */
    .stSpinner > div {
        border-top-color: var(--primary-color) !important;
    }
    
    /* Section headers in tabs */
    .section-header {
        color: var(--primary-color);
        font-size: 1.5rem;
        font-weight: 600;
        margin-bottom: 1rem;
        padding-bottom: 0.5rem;
        border-bottom: 2px solid var(--border-color);
    }
</style>
""", unsafe_allow_html=True)

# --- Session State Initialization ---
if "templates" not in st.session_state:
    st.session_state.templates = {}
if "generated_content" not in st.session_state:
    st.session_state.generated_content = None 
if "template_bytes_for_download" not in st.session_state:
    st.session_state.template_bytes_for_download = None 
if "file_name_info" not in st.session_state:
    st.session_state.file_name_info = None 
if "final_docx_bytes" not in st.session_state:
    st.session_state.final_docx_bytes = None
if "final_file_name" not in st.session_state:
    st.session_state.final_file_name = None

# --- Sidebar: Template Library Manager ---
with st.sidebar:
    st.markdown("## Template Library")
    
    st.divider()
    
    with st.form("new_template_form", clear_on_submit=True):
        st.markdown("### Add New Template")
        new_template_name = st.text_input("Template Name", placeholder="e.g., 'UCSF IRB'")
        
        st.markdown("**1. DOCX Structure**")
        new_template_file = st.file_uploader(
            "Required section headings", 
            type="docx",
            help="The .docx file with the section headings.",
            label_visibility="collapsed"
        )
        
        st.markdown("**2. Approved Language (Optional)**")
        new_language_file = st.file_uploader(
            "PDF, DOCX, TXT with boilerplate",
            type=["pdf", "docx", "txt", "md"],
            help="Optional: The doc with IRB-approved 'boilerplate' text.",
            label_visibility="collapsed"
        )
        
        submitted = st.form_submit_button("Save Template", use_container_width=True)
        
        if submitted:
            if new_template_name and new_template_file:
                if new_template_name in st.session_state.templates:
                    st.warning(f"Template '{new_template_name}' already exists. Overwriting.")
                
                template_bytes = new_template_file.getvalue()
                template_data = {
                    "template_filename": new_template_file.name,
                    "template_bytes": template_bytes
                }
                
                if new_language_file:
                    language_bytes = new_language_file.getvalue()
                    template_data["language_filename"] = new_language_file.name
                    template_data["language_bytes"] = language_bytes
                
                st.session_state.templates[new_template_name] = template_data
                st.success(f"✓ Saved: {new_template_name}")
            else:
                st.error("Please provide Template Name and DOCX file.")
    
    if st.session_state.templates:
        st.divider()
        st.markdown("### Current Library")
        for t_name, data in st.session_state.templates.items():
            structure = "plain language ICF template.docx"
            language = ""
            if data.get("language_filename"):
                language = f"<br><span style='color: #2DD4BF;'>Language:</span> {data['language_filename']}"
            
            st.markdown(f"""
            <div class="template-card">
                <div class="template-card-title">{t_name}</div>
                <div class="template-card-subtitle">
                    <span style='color: #737373;'>Structure:</span> {structure}{language}
                </div>
            </div>
            """, unsafe_allow_html=True)


# --- REAL API FUNCTIONS (keeping all original functions) ---

def extract_text_from_pdf(pdf_bytes: bytes) -> str:
    """Extract text from PDF using PyPDF2"""
    try:
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(pdf_bytes))
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        st.error(f"Error extracting PDF text: {e}")
        return ""

def extract_text_from_docx(docx_bytes: bytes) -> str:
    """Extract text from DOCX"""
    try:
        doc = Document(io.BytesIO(docx_bytes))
        text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        return text
    except Exception as e:
        st.error(f"Error extracting DOCX text: {e}")
        return ""

def call_claude_api(prompt: str, system_prompt: str = "", max_tokens: int = 4096) -> str:
    """Call Claude API for text generation"""
    claude_api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not claude_api_key:
        return "" 
    
    try:
        client = anthropic.Anthropic(api_key=claude_api_key)
        
        message = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=max_tokens,
            system=system_prompt if system_prompt else anthropic.NOT_GIVEN,
            messages=[
                {"role": "user", "content": prompt}
            ]
        )
        
        return message.content[0].text
    except Exception as e:
        st.error(f"Claude API Error: {e}")
        return ""

def parse_protocol_with_claude(protocol_text: str, file_type: str) -> Optional[Dict]:
    """Use Claude to extract structured data from protocol"""
    
    system_prompt = """You are a clinical research expert. Extract key information from research protocols and return it as JSON.
    
    Extract these fields:
    - title: Study title
    - purpose: Primary objective
    - procedures: List of study procedures
    - risks: List of risks (each with 'name' and 'details')
    - benefits: Potential benefits
    - alternatives: Alternative treatments
    - compensation: Payment details
    - confidentiality: Privacy/confidentiality information
    
    Return ONLY valid JSON, no additional text."""
    
    prompt = f"""Analyze this clinical research protocol and extract the required information:

{protocol_text[:15000]}  

Return the information as a JSON object with the structure specified in the system prompt."""
    
    try:
        response = call_claude_api(prompt, system_prompt, max_tokens=4096)
        
        if not response:
            return None
        
        if "```json" in response:
            json_str = response.split("```json")[1].split("```")[0].strip()
        elif "```" in response:
            json_str = response.split("```")[1].split("```")[0].strip()
        else:
            json_str = response.strip()
        
        return json.loads(json_str)
    except json.JSONDecodeError as e:
        st.error(f"Error parsing Claude response as JSON: {e}")
        return None
    except Exception as e:
        st.error(f"Error processing protocol: {e}")
        return None

def parse_approved_language_with_claude(language_text: str, filename: str) -> Dict:
    """Use Claude to extract approved language snippets"""
    
    system_prompt = """You are extracting IRB-approved language templates. Return as JSON with these keys:
    - risks_intro: Introduction to risks section
    - confidentiality: HIPAA/privacy language
    - payment: Compensation language
    - voluntary: Voluntary participation language
    
    Return ONLY valid JSON."""
    
    prompt = f"""Extract IRB-approved language templates from this document:

{language_text[:10000]}

Return as JSON with the keys specified."""
    
    try:
        response = call_claude_api(prompt, system_prompt, max_tokens=2048)
        
        if "```json" in response:
            json_str = response.split("```json")[1].split("```")[0].strip()
        elif "```" in response:
            json_str = response.split("```")[1].split("```")[0].strip()
        else:
            json_str = response.strip()
        
        return json.loads(json_str)
    except Exception as e:
        st.warning(f"Could not parse approved language: {e}")
        return {}

def generate_icf_content_with_claude(
    protocol_data: Dict, 
    approved_language: Dict, 
    template_headings: List[str]
) -> Dict[str, str]:
    """Use Claude to generate ICF content for each section"""
    
    system_prompt = """You are a clinical research professional creating an Informed Consent Form (ICF). 
    
    Generate formal, compliant ICF text at a collegiate reading level. Use approved language when provided.
    Be thorough, accurate, and maintain appropriate medical/legal terminology.
    
    Return ONLY the generated text for each section, no additional formatting or comments."""
    
    generated_content = {}
    
    for heading in template_headings:
        prompt = f"""Generate content for this ICF section: "{heading}"

Protocol Information:
{json.dumps(protocol_data, indent=2)}

Approved Language Templates:
{json.dumps(approved_language, indent=2)}

Generate formal, compliant text for the "{heading}" section. Use approved language where applicable.
If this section relates to risks, use the risks_intro template as a starting point.
If this section relates to confidentiality, use the confidentiality template.

Return ONLY the paragraph text for this section, no headers or additional formatting."""
        
        try:
            content = call_claude_api(prompt, system_prompt, max_tokens=1500)
            generated_content[heading] = content if content else f"Content for {heading} pending review."
        except Exception as e:
            st.warning(f"Error generating content for {heading}: {e}")
            generated_content[heading] = f"[Content for {heading} requires manual entry]"
    
    return generated_content

def analyze_icf_with_claude(
    protocol_data: Dict,
    approved_language: Dict,
    icf_text: str
) -> Dict:
    """Use Claude to analyze ICF for gaps and issues"""
    
    system_prompt = """You are a clinical research compliance expert reviewing Informed Consent Forms.
    
    Analyze the ICF for:
    1. Readability (estimate grade level)
    2. Gaps compared to protocol
    3. Missing approved language
    4. Suggestions for improvement
    
    Return as JSON with keys: overall_score, gaps (list of strings), suggestions (list of objects with 'section', 'issue', 'suggestion')"""
    
    prompt = f"""Analyze this ICF for compliance and completeness:

PROTOCOL DATA:
{json.dumps(protocol_data, indent=2)}

APPROVED LANGUAGE REQUIREMENTS:
{json.dumps(approved_language, indent=2)}

ICF TEXT:
{icf_text[:20000]}

Identify:
1. Estimated reading grade level (overall_score)
2. Missing protocol information (gaps list)
3. Missing required language (gaps list)
4. Specific improvement suggestions

Return as JSON."""
    
    try:
        response = call_claude_api(prompt, system_prompt, max_tokens=3000)
        
        if "```json" in response:
            json_str = response.split("```json")[1].split("```")[0].strip()
        elif "```" in response:
            json_str = response.split("```")[1].split("```")[0].strip()
        else:
            json_str = response.strip()
        
        analysis = json.loads(json_str)
        
        sections = parse_docx_sections_from_text(icf_text)
        readability_report = []
        for section, text in sections.items():
            if len(text) > 50:
                score = textstat.flesch_kincaid_grade(text)
                readability_report.append({"section": section, "score": score})
        
        analysis["readability_report"] = readability_report
        return analysis
        
    except Exception as e:
        st.error(f"Error analyzing ICF: {e}")
        return {"overall_score": 0, "gaps": [], "suggestions": [], "readability_report": []}

def summarize_icf_with_claude(sections: Dict[str, str], target_level: str) -> Dict[str, str]:
    """Use Claude to create patient-friendly summaries"""
    
    grade_level_guidance = {
        "6th Grade": "Use very simple language. Short sentences. Common everyday words. Explain medical terms in plain English.",
        "8th Grade": "Use clear, straightforward language. Moderate sentence length. Define technical terms when used."
    }
    
    system_prompt = f"""You are creating patient-friendly summaries of medical research consent forms.
    
    Target reading level: {target_level}
    Guidance: {grade_level_guidance[target_level]}
    
    Make the content accessible and easy to understand while maintaining accuracy."""
    
    summaries = {}
    
    for heading, text in sections.items():
        if len(text.strip()) < 20:
            summaries[heading] = text
            continue
            
        prompt = f"""Summarize this ICF section for a patient at a {target_level} reading level:

Section: {heading}

Original text:
{text[:2000]}

Create a clear, simple summary that a {target_level} student could understand."""
        
        try:
            summary = call_claude_api(prompt, system_prompt, max_tokens=800)
            summaries[heading] = summary if summary else text
        except Exception as e:
            st.warning(f"Error summarizing {heading}: {e}")
            summaries[heading] = text
    
    return summaries

def parse_docx_sections_from_text(text: str) -> Dict[str, str]:
    """Simple section parser for text"""
    sections = {}
    lines = text.split('\n')
    current_heading = "Introduction"
    current_text = []
    
    for line in lines:
        if line.strip() and (line.isupper() or line.endswith(':')):
            if current_text:
                sections[current_heading] = '\n'.join(current_text)
            current_heading = line.strip()
            current_text = []
        elif line.strip():
            current_text.append(line.strip())
    
    if current_text:
        sections[current_heading] = '\n'.join(current_text)
    
    return sections if sections else {"Full Document": text}

def parse_template_headings(docx_bytes):
    try:
        doc = Document(io.BytesIO(docx_bytes))
        headings = []
        for p in doc.paragraphs:
            if p.style and p.style.name.startswith(('Heading 1', 'Heading 2', 'Heading 3')):
                if p.text.strip():
                    headings.append(p.text.strip())
        if not headings:
            for p in doc.paragraphs:
                if p.text.strip() and any(r.bold for r in p.runs):
                    headings.append(p.text.strip())
        if not headings:
            return ["Introduction", "Purpose", "Procedures", "Risks", "Benefits", "Alternatives", "Compensation", "Confidentiality"], doc
        return list(dict.fromkeys(headings)), doc
    except Exception as e:
        st.error(f"Error parsing DOCX file: {e}")
        return None, None

def populate_docx_template(template_doc_bytes, generated_content):
    """Populate DOCX template with generated content"""
    try:
        doc = Document(io.BytesIO(template_doc_bytes))
        paras = doc.paragraphs
        target_headings = set(generated_content.keys())
        
        all_heading_texts = set()
        for p in paras:
            p_text = p.text.strip()
            if not p_text:
                continue
            style_name = p.style.name if p.style else 'Normal'
            is_heading_style = style_name.startswith(('Heading 1', 'Heading 2', 'Heading 3'))
            is_bold = any(r.bold for r in p.runs)
            if is_heading_style or (is_bold and not style_name.startswith('Heading')):
                all_heading_texts.add(p_text)

        i = 0
        while i < len(paras):
            p = paras[i]
            p_text = p.text.strip()

            if p_text in target_headings:
                current_heading = p_text
                j = i + 1
                while j < len(paras):
                    next_p_text = paras[j].text.strip()
                    if next_p_text in all_heading_texts:
                        break
                    j += 1
                
                for k in range(j - 1, i, -1):
                    p_to_delete = paras[k]
                    p_to_delete._element.getparent().remove(p_to_delete._element)
                
                text_to_add = generated_content[current_heading]
                new_p = doc.add_paragraph(text_to_add, style='Normal')
                new_p.paragraph_format.space_before = Pt(6)
                new_p.paragraph_format.space_after = Pt(12)
                
                p_to_move = new_p._p
                p_after = p._p
                p_after.addnext(p_to_move)
                i = j
            else:
                i += 1
        
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream.getvalue()
    except Exception as e:
        st.error(f"Error populating DOCX template: {e}")
        return None

def parse_docx_sections(docx_bytes):
    try:
        doc = Document(io.BytesIO(docx_bytes))
        sections = {}
        current_heading = "Introduction"
        current_text = []
        for p in doc.paragraphs:
            style_name = p.style.name if p.style else 'Normal'
            is_heading = style_name.startswith(('Heading 1', 'Heading 2', 'Heading 3'))
            is_bold = any(r.bold for r in p.runs)
            
            if p.text.strip() and (is_heading or (is_bold and not style_name.startswith('Heading'))):
                if current_text:
                    sections[current_heading] = "\n".join(current_text)
                current_heading = p.text.strip()
                current_text = []
            elif p.text.strip():
                current_text.append(p.text.strip())
        
        if current_text:
            sections[current_heading] = "\n".join(current_text)
            
        if not sections:
            return {"error": "Could not read any text sections from the uploaded DOCX."}
        return sections
    except Exception as e:
        return {"error": f"Error parsing ICF DOCX: {e}"}


# --- Main Application UI ---
st.markdown('<h1><span>ICF-GPT:</span> Clinical Consent AI</h1>', unsafe_allow_html=True)
st.markdown("Generate compliant ICFs from protocols using **Claude 4 Sonnet API** for real-time analysis.")

# Check API configuration
if not os.environ.get("ANTHROPIC_API_KEY"):
    st.error("⚠️ **Anthropic API Key Not Found.** Set the `ANTHROPIC_API_KEY` environment variable to enable AI functionality.")

# Custom styled tabs
tab_generate, tab_review, tab_summarize = st.tabs([
    "✨ Generate New ICF", 
    "🔍 Review & Analyze", 
    "👤 Patient Summarizer"
])

# --- TAB 1: GENERATION MODE ---
with tab_generate:
    st.markdown('<div class="section-header">Generate New ICF</div>', unsafe_allow_html=True)
    st.markdown("Upload your protocol and select a template to generate an AI-powered draft.")
    
    with st.form("generation_form"):
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("#### 📄 Upload Protocol")
            protocol_file = st.file_uploader(
                "Upload Protocol Document", 
                type=["pdf", "docx", "json", "xml"], 
                help="The source protocol document.",
                label_visibility="collapsed"
            )
        
        with col2:
            st.markdown("#### 📋 Select Template")
            template_options = list(st.session_state.templates.keys())
            selected_template_name = None
            
            if not template_options:
                st.warning("No saved templates. Add one in the sidebar.")
            else:
                selected_template_name = st.selectbox(
                    "Select IRB Template",
                    options=template_options,
                    help="Choose from your saved templates.",
                    label_visibility="collapsed"
                )
        
        submitted = st.form_submit_button("✨ Generate AI Draft", type="primary", use_container_width=True)

    if submitted:
        st.session_state.final_docx_bytes = None
        st.session_state.final_file_name = None
        
        template_info = None
        if selected_template_name:
            template_info = st.session_state.templates[selected_template_name]
        
        if protocol_file is not None and template_info is not None:
            if not os.environ.get("ANTHROPIC_API_KEY"):
                st.error("Please set the `ANTHROPIC_API_KEY` environment variable.")
            else:
                with st.spinner("🤖 Claude is analyzing your protocol..."):
                    protocol_bytes = protocol_file.getvalue()
                    file_ext = protocol_file.name.split('.')[-1].lower()
                    
                    if file_ext == "pdf":
                        protocol_text = extract_text_from_pdf(protocol_bytes)
                    elif file_ext == "docx":
                        protocol_text = extract_text_from_docx(protocol_bytes)
                    elif file_ext in ["json", "xml"]:
                        protocol_text = protocol_bytes.decode('utf-8')
                    else:
                        st.error(f"Unsupported file type: {file_ext}")
                        protocol_text = None
                    
                    if protocol_text:
                        protocol_data = parse_protocol_with_claude(protocol_text, file_ext)
                        
                        if protocol_data:
                            template_bytes = template_info["template_bytes"]
                            template_headings, _ = parse_template_headings(template_bytes)
                            
                            approved_language_data = {}
                            if "language_bytes" in template_info:
                                lang_bytes = template_info["language_bytes"]
                                lang_ext = template_info["language_filename"].split('.')[-1].lower()
                                
                                if lang_ext == "pdf":
                                    lang_text = extract_text_from_pdf(lang_bytes)
                                elif lang_ext == "docx":
                                    lang_text = extract_text_from_docx(lang_bytes)
                                else:
                                    lang_text = lang_bytes.decode('utf-8')
                                
                                approved_language_data = parse_approved_language_with_claude(
                                    lang_text, template_info["language_filename"]
                                )
                            
                            with st.spinner("✍️ Claude is drafting your ICF sections..."):
                                generated_content = generate_icf_content_with_claude(
                                    protocol_data, 
                                    approved_language_data, 
                                    template_headings
                                )
                                
                                st.session_state.generated_content = generated_content
                                st.session_state.template_bytes_for_download = template_bytes
                                st.session_state.file_name_info = f"Generated_ICF_{selected_template_name.replace(' ', '_')}.docx"
                                st.success("✅ AI Draft Complete!")
        else:
            st.error("🚫 Please upload a Protocol and select a Template.")

    # Editor UI
    if st.session_state.generated_content:
        st.divider()
        st.markdown('<div class="section-header">Review & Finalize</div>', unsafe_allow_html=True)
        st.info("Review and edit the AI-generated content before downloading.")
        
        with st.form("editor_form"):
            final_content_map = {}
            for heading, text in st.session_state.generated_content.items():
                st.markdown(f"#### {heading}")
                edited_text = st.text_area(
                    f"Edit {heading}", 
                    value=text, 
                    height=200, 
                    key=f"edit_{heading}",
                    label_visibility="collapsed"
                )
                final_content_map[heading] = edited_text

            download_submitted = st.form_submit_button("✅ Finalize Document", type="primary", use_container_width=True)

            if download_submitted:
                with st.spinner("Creating final DOCX..."):
                    template_bytes = st.session_state.template_bytes_for_download
                    output_docx_bytes = populate_docx_template(template_bytes, final_content_map)
                    
                    if output_docx_bytes:
                        st.session_state.final_docx_bytes = output_docx_bytes
                        st.session_state.final_file_name = st.session_state.file_name_info
                        st.session_state.generated_content = None
                        st.session_state.template_bytes_for_download = None
                        st.session_state.file_name_info = None
                        st.success("✅ Document ready for download!")
                        st.rerun()
                        
    # Download button
    if st.session_state.final_docx_bytes:
        st.download_button(
            label="📥 Download Final ICF (.docx)",
            data=st.session_state.final_docx_bytes,
            file_name=st.session_state.final_file_name,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
            on_click=lambda: st.session_state.update(final_docx_bytes=None, final_file_name=None)
        )
        st.info("Click to download. Button will disappear after clicking.")


# --- TAB 2: REVIEW MODE ---
with tab_review:
    st.markdown('<div class="section-header">Review & Analyze ICF</div>', unsafe_allow_html=True)
    st.markdown("Upload your protocol, template, and existing ICF for comprehensive AI analysis.")
    
    with st.form("review_form"):
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("#### 📄 Source Documents")
            protocol_file_review = st.file_uploader(
                "Upload Protocol", 
                type=["pdf", "docx", "json", "xml"], 
                key="protocol_rev",
                label_visibility="collapsed"
            )
            
            template_options_rev = list(st.session_state.templates.keys())
            selected_template_name_rev = None
            if template_options_rev:
                selected_template_name_rev = st.selectbox(
                    "Select IRB Template",
                    options=template_options_rev,
                    key="template_rev"
                )

        with col2:
            st.markdown("#### 📋 ICF to Review")
            existing_icf_docx = st.file_uploader(
                "Upload Existing ICF (DOCX)", 
                type="docx", 
                key="docx_rev",
                label_visibility="collapsed"
            )

        review_submitted = st.form_submit_button("🔍 Analyze with Claude", type="primary", use_container_width=True)

    if review_submitted:
        template_info_rev = None
        if selected_template_name_rev:
            template_info_rev = st.session_state.templates[selected_template_name_rev]
            
        if protocol_file_review and existing_icf_docx and template_info_rev:
            if not os.environ.get("ANTHROPIC_API_KEY"):
                st.error("Please set the `ANTHROPIC_API_KEY` environment variable.")
            else:
                with st.spinner("🤖 Claude is analyzing documents..."):
                    protocol_bytes = protocol_file_review.getvalue()
                    file_ext = protocol_file_review.name.split('.')[-1].lower()
                    
                    if file_ext == "pdf":
                        protocol_text = extract_text_from_pdf(protocol_bytes)
                    elif file_ext == "docx":
                        protocol_text = extract_text_from_docx(protocol_bytes)
                    else:
                        protocol_text = protocol_bytes.decode('utf-8')
                    
                    protocol_data = parse_protocol_with_claude(protocol_text, file_ext)

                    icf_bytes = existing_icf_docx.getvalue()
                    icf_text = extract_text_from_docx(icf_bytes)

                    approved_language_data = {}
                    if "language_bytes" in template_info_rev:
                        lang_bytes = template_info_rev["language_bytes"]
                        lang_ext = template_info_rev["language_filename"].split('.')[-1].lower()
                        
                        if lang_ext == "pdf":
                            lang_text = extract_text_from_pdf(lang_bytes)
                        elif lang_ext == "docx":
                            lang_text = extract_text_from_docx(lang_bytes)
                        else:
                            lang_text = lang_bytes.decode('utf-8')
                        
                        approved_language_data = parse_approved_language_with_claude(
                            lang_text, template_info_rev["language_filename"]
                        )
                    
                    analysis_result = analyze_icf_with_claude(
                        protocol_data, 
                        approved_language_data, 
                        icf_text
                    )
                    
                    st.session_state.analysis_result = analysis_result
                    st.success("✅ Analysis Complete!")
        else:
            st.error("🚫 Please upload all required documents.")

    if "analysis_result" in st.session_state and st.session_state.analysis_result:
        analysis_result = st.session_state.analysis_result
        st.divider()
        st.markdown('<div class="section-header">Analysis Report</div>', unsafe_allow_html=True)

        col_a1, col_a2 = st.columns(2)
        with col_a1:
            st.metric("AI Estimated Grade Level", f"{analysis_result.get('overall_score', 'N/A')}")
        with col_a2:
            icf_text_check = extract_text_from_docx(icf_bytes) if 'icf_bytes' in locals() else ""
            fk_score = textstat.flesch_kincaid_grade(icf_text_check) if icf_text_check and len(icf_text_check) > 50 else "N/A"
            st.metric("Flesch-Kincaid Grade", f"{fk_score}")

        st.markdown("#### 🔍 Gaps & Missing Information")
        if analysis_result.get("gaps"):
            for gap in analysis_result["gaps"]:
                st.warning(f"• {gap}")
        else:
            st.success("✓ No significant gaps identified")
            
        st.markdown("#### 💡 Improvement Suggestions")
        suggestions = analysis_result.get("suggestions", [])
        if suggestions:
            for s in suggestions:
                with st.expander(f"**{s.get('section', 'General')}**"):
                    st.write(f"**Issue:** {s.get('issue', 'N/A')}")
                    st.write(f"**Suggestion:** {s.get('suggestion', 'N/A')}")
        else:
            st.success("✓ No critical improvements needed")

        st.markdown("#### 📊 Section Readability Breakdown")
        readability_report = analysis_result.get("readability_report", [])
        if readability_report:
            df = pd.DataFrame(readability_report)
            st.dataframe(
                df.style.highlight_max(axis=0, subset=['score'], color='#ff9999')
                       .highlight_min(axis=0, subset=['score'], color='#ccffcc'), 
                use_container_width=True
            )


# --- TAB 3: SUMMARIZER MODE ---
with tab_summarize:
    st.markdown('<div class="section-header">Patient Summarizer</div>', unsafe_allow_html=True)
    st.markdown("Improve readability by generating plain-language summaries of an ICF.")
    
    st.markdown("### Step 1: Upload and Configure")

    with st.form("summarize_form"):
        st.markdown("#### Upload Completed ICF")
        icf_file_summarize = st.file_uploader(
            "Upload ICF Document", 
            type="docx", 
            key="icf_sum",
            label_visibility="collapsed"
        )
        
        # Fixed to 6th grade level
        target_grade = "6th Grade"
        
        summarize_submitted = st.form_submit_button("📝 Generate Summaries", type="primary", use_container_width=True)
        
    if summarize_submitted:
        if icf_file_summarize:
            if not os.environ.get("ANTHROPIC_API_KEY"):
                st.error("Please set the `ANTHROPIC_API_KEY` environment variable.")
            else:
                with st.spinner(f"✍️ Claude is summarizing to {target_grade} level..."):
                    icf_bytes = icf_file_summarize.getvalue()
                    icf_sections = parse_docx_sections(icf_bytes)
                    
                    if "error" in icf_sections:
                        st.error(icf_sections["error"])
                    else:
                        summaries = summarize_icf_with_claude(icf_sections, target_grade)
                        st.session_state.summaries = summaries
                        st.session_state.file_name_summarize = f"Summary_{icf_file_summarize.name}"
                        st.success("✅ Summary Complete!")
        else:
            st.error("🚫 Please upload an ICF document.")

    if "summaries" in st.session_state and st.session_state.summaries:
        st.divider()
        st.markdown('<div class="section-header">Patient Summary Review</div>', unsafe_allow_html=True)
        
        summarized_doc = Document()
        
        for heading, summary_text in st.session_state.summaries.items():
            h = summarized_doc.add_heading(heading, level=1)
            h.runs[0].bold = True
            h.runs[0].font.size = Pt(14)
            p = summarized_doc.add_paragraph(summary_text, style='Normal')
            p.paragraph_format.space_after = Pt(12)
            
            with st.expander(f"**{heading}**", expanded=True):
                st.write(summary_text)

        file_stream = io.BytesIO()
        summarized_doc.save(file_stream)
        file_stream.seek(0)
        summary_docx_bytes = file_stream.getvalue()
        
        st.divider()
        st.download_button(
            label="📥 Download Patient Summary (.docx)",
            data=summary_docx_bytes,
            file_name=f"Patient_Summary_{st.session_state.file_name_summarize}",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
